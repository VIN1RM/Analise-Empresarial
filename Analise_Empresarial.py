import tkinter as tk
from tkinter import messagebox
import pandas as pd
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import seaborn as sns
import os

# Função para solicitar o nome da empresa
def solicitar_nome_empresa():
    root = tk.Tk()
    root.title("Nome da Empresa")
    root.geometry("400x150")
    
    tk.Label(root, text="Informe o nome da empresa para análise:", font=("Arial", 12)).pack(pady=10)
    entry_nome = tk.Entry(root, font=("Arial", 12))
    entry_nome.pack(pady=5)
    
    def confirmar():
        nome_empresa = entry_nome.get().strip()
        if not nome_empresa:
            messagebox.showerror("Erro", "Por favor, insira o nome da empresa.")
        else:
            root.destroy()
            coletar_dados_via_texto(nome_empresa)
    
    btn_confirmar = tk.Button(root, text="Confirmar", command=confirmar, font=("Arial", 12), bg="green", fg="white")
    btn_confirmar.pack(pady=10)
    
    root.mainloop()

# Função para coletar os dados via texto colado do Excel
def coletar_dados_via_texto(nome_empresa):
    root = tk.Tk()
    root.title("Inserção de Dados")
    root.geometry("800x500")

    tk.Label(root, text="Cole os dados copiados do Excel abaixo e clique em 'Confirmar'", font=("Arial", 12)).pack()

    text_area = tk.Text(root, wrap=tk.WORD, font=("Arial", 10))
    text_area.pack(expand=True, fill="both", padx=10, pady=10)

    def processar_texto():
        conteudo = text_area.get("1.0", tk.END).strip()
        linhas = conteudo.split("\n")
        dados = [linha.split("\t") for linha in linhas]

        if len(dados) > 1:
            colunas = dados[0]
            valores = dados[1:]
            df = pd.DataFrame(valores, columns=colunas)
            df.replace("%", "", regex=True, inplace=True)
            df = df.apply(pd.to_numeric, errors='ignore')
            root.destroy()
            gerar_relatorio(df, nome_empresa)
        else:
            messagebox.showerror("Erro", "Nenhum dado válido foi inserido.")

    btn_confirmar = tk.Button(root, text="Confirmar", command=processar_texto, font=("Arial", 12), bg="green", fg="white")
    btn_confirmar.pack(pady=10)

    root.mainloop()

# Função para gerar gráficos automaticamente
def gerar_graficos(df, destino):
    colunas_excluir = ["Nota", "Média", "Média Geral"]
    disciplinas = [col for col in df.columns[1:] if col not in colunas_excluir]

    # Gráfico Pizza (necessidade de reforço por disciplina)
    reforco_disciplina = (df[disciplinas].apply(pd.to_numeric, errors='coerce') < 80).sum()
    reforco_disciplina = reforco_disciplina[reforco_disciplina > 0]
    plt.figure(figsize=(8, 8))
    plt.pie(reforco_disciplina, labels=reforco_disciplina.index, autopct='%1.1f%%', 
            colors=sns.color_palette("coolwarm", len(reforco_disciplina)),
            wedgeprops={'edgecolor': 'white'}, pctdistance=0.85, labeldistance=1.1)
    plt.ylabel("")
    plt.title("Proporção de Disciplinas com Necessidade de Reforço")
    plt.savefig(os.path.join(destino, "grafico_pizza.png"), bbox_inches='tight')
    plt.close()

    # Gráfico de Barras (média das disciplinas)
    plt.figure(figsize=(10, 6))
    medias = df[disciplinas].apply(pd.to_numeric, errors='coerce').mean()
    cores = ['#3498db' if media >= 80 else '#e74c3c' for media in medias]
    sns.barplot(x=medias.index, y=medias.values, palette=cores)
    plt.axhline(80, color='black', linestyle='--', label='Linha de Corte 80%')
    plt.xticks(rotation=45, ha='right')
    plt.ylabel('Média das Notas (%)')
    plt.xlabel('Disciplinas')
    plt.title("Média das Notas por Disciplina")
    plt.legend()
    plt.tight_layout()
    plt.savefig(os.path.join(destino, "grafico_barras.png"))
    plt.close()

    # Gráficos Individuais para cada aluno
    for _, aluno in df.iterrows():
        plt.figure(figsize=(10,6))
        notas_aluno = aluno[disciplinas].apply(pd.to_numeric, errors='coerce')
        cores_individual = ['#3498db' if nota >= 80 else '#e74c3c' for nota in notas_aluno]
        sns.barplot(x=disciplinas, y=notas_aluno, palette=cores_individual)
        plt.axhline(80, color='black', linestyle='--', label='Linha de Corte 80%')
        plt.xticks(rotation=45, ha='right')
        plt.ylim(0, 100)
        plt.ylabel('Nota (%)')
        plt.title(f"Desempenho de {aluno[df.columns[0]]}")
        plt.legend()
        plt.tight_layout()
        plt.savefig(os.path.join(destino, f"grafico_{aluno[df.columns[0]]}.png"))
        plt.close()

    # --- Gráficos: Relatório de Desempenho por Disciplina ---
    for disciplina in disciplinas:
        plt.figure(figsize=(10,6))
        notas = pd.to_numeric(df[disciplina], errors='coerce')
        bins = range(0, 101, 10)
        # rwidth ajustado para deixar as colunas mais finas
        plt.hist(notas, bins=bins, color="#3498db", edgecolor="black", rwidth=0.5)
        plt.xlabel('Nota (%)')
        plt.ylabel('Quantidade de colaboradores')
        plt.title(f"Distribuição de Notas na Disciplina: {disciplina}")
        plt.axvline(80, color='red', linestyle='--', label='Linha de Corte 80%')
        plt.legend()
        plt.tight_layout()
        plt.savefig(os.path.join(destino, f"grafico_disciplina_{disciplina}.png"))
        plt.close()

# Função para gerar relatório
def gerar_relatorio(df, nome_empresa):
    # Cria uma pasta na Área de Trabalho com o nome "Documentação <Nome da Empresa>"
    desktop = os.path.join(os.path.expanduser("~"), "Desktop")
    pasta_destino = os.path.join(desktop, f"Documentação {nome_empresa}")
    os.makedirs(pasta_destino, exist_ok=True)
    
    gerar_graficos(df, pasta_destino)

    colunas_excluir = ["Nota", "Média", "Média Geral"]
    disciplinas = [col for col in df.columns[1:] if col not in colunas_excluir]

    doc = Document()

    doc.add_heading('ANÁLISE DE ALUNOS E DISCIPLINAS PARA REFORÇO', level=1)
    doc.add_paragraph("""
    Esta análise identifica alunos com desempenho abaixo de 80 para reforço nas disciplinas específicas.
    """)

    contador = 1
    for _, row in df.iterrows():
        aluno = row[df.columns[0]]
        disciplinas_reforco = [d for d in disciplinas if pd.to_numeric(row[d], errors='coerce') < 80]
        if disciplinas_reforco:
            doc.add_paragraph(f"{contador}. {aluno}", style='Heading 2')
            doc.add_paragraph(f"Disciplinas: {', '.join(disciplinas_reforco)}")
            contador += 1

    doc.add_page_break()
    doc.add_heading("Gráficos Gerais", level=1)
    doc.add_picture(os.path.join(pasta_destino, "grafico_pizza.png"), width=Inches(5))
    doc.add_picture(os.path.join(pasta_destino, "grafico_barras.png"), width=Inches(5))

    doc.add_page_break()
    doc.add_heading('RELATÓRIO DE DESEMPENHO INDIVIDUAL', level=1)
    for _, row in df.iterrows():
        aluno = row[df.columns[0]]
        grafico_individual = os.path.join(pasta_destino, f"grafico_{aluno}.png")
        doc.add_page_break()
        doc.add_heading(f"Desempenho de {aluno}", level=2)
        doc.add_picture(grafico_individual, width=Inches(5))

    # Seção: Gráficos por Disciplina com Lista de Colaboradores
    doc.add_page_break()
    doc.add_heading('RELATÓRIO DE DESEMPENHO POR DISCIPLINA', level=1)
    for disciplina in disciplinas:
        grafico_disciplina = os.path.join(pasta_destino, f"grafico_disciplina_{disciplina}.png")
        doc.add_heading(f"Desempenho na Disciplina: {disciplina}", level=2)
        doc.add_picture(grafico_disciplina, width=Inches(5))
        # Ordena os colaboradores pela nota na disciplina (ordem decrescente)
        df_sorted = df.sort_values(by=disciplina, ascending=False)
        nomes_ordenados = df_sorted[df.columns[0]].tolist()
        nomes_str = ", ".join(nomes_ordenados)
        doc.add_paragraph("Ordem dos colaboradores: " + nomes_str)

    relatorio_path = os.path.join(pasta_destino, "relatorio.docx")
    doc.save(relatorio_path)
    messagebox.showinfo("Sucesso", f"Relatório criado com sucesso em:\n{relatorio_path}")

# Inicia a aplicação
def iniciar_aplicacao():
    root = tk.Tk()
    root.withdraw()
    resposta = messagebox.askquestion("Iniciar", "Deseja iniciar uma nova análise e documentação?")
    if resposta == 'yes':
        solicitar_nome_empresa()
    else:
        messagebox.showinfo("Encerrado", "O programa foi encerrado.")

if __name__ == "__main__":
    iniciar_aplicacao()